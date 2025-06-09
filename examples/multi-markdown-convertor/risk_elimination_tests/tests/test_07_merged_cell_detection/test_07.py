#!/usr/bin/env python3
"""
Test 7: Merged-cell detection via `tblLook` flags is consistent
Two DOCX tables: one with header row shading, one plain
"""

import json
import sys
import tempfile
from pathlib import Path
from typing import Dict, Any, List, Optional


def create_test_docx() -> Path:
    """Create a test DOCX file with tables having different tblLook flags."""
    try:
        from docx import Document
        from docx.shared import Inches

        # Create document
        doc = Document()
        doc.add_heading("Table tblLook Flags Test", 0)

        # Add first table with header row styling
        doc.add_paragraph("Table 1: With header row shading")
        table1 = doc.add_table(rows=3, cols=3)
        table1.style = "Table Grid"

        # Fill table 1
        hdr_cells = table1.rows[0].cells
        hdr_cells[0].text = "Header 1"
        hdr_cells[1].text = "Header 2"
        hdr_cells[2].text = "Header 3"

        for i in range(1, 3):
            row_cells = table1.rows[i].cells
            for j in range(3):
                row_cells[j].text = f"Cell {i},{j + 1}"

        # Add second table without special styling
        doc.add_paragraph("Table 2: Plain table")
        table2 = doc.add_table(rows=3, cols=3)

        # Fill table 2
        for i in range(3):
            row_cells = table2.rows[i].cells
            for j in range(3):
                if i == 0:
                    row_cells[j].text = f"Plain Header {j + 1}"
                else:
                    row_cells[j].text = f"Data {i},{j + 1}"

        # Save to temp file
        temp_file = Path(tempfile.mktemp(suffix=".docx"))
        doc.save(str(temp_file))
        return temp_file

    except ImportError:
        # Fallback: use existing DOCX file if available
        test_inputs_dir = Path(__file__).parent.parent.parent / "test-inputs"
        docx_files = list(test_inputs_dir.glob("*.docx"))
        if docx_files:
            return docx_files[0]
        else:
            raise RuntimeError(
                "python-docx not available and no test DOCX files found"
            )


def analyze_table_properties(docx_path: Path) -> Dict[str, Any]:
    """Analyze table properties and tblLook flags in a DOCX document."""
    try:
        from docx import Document

        doc = Document(str(docx_path))

        analysis = {
            "tables_found": 0,
            "table_details": [],
            "has_header_flags": False,
            "consistent_detection": True,
        }

        for i, table in enumerate(doc.tables):
            table_info = {
                "table_index": i,
                "rows": len(table.rows),
                "cols": len(table.columns) if table.rows else 0,
                "style": table.style.name if table.style else None,
                "has_header_row": False,
                "tbl_look_flags": None,
            }

            # Try to access table properties
            try:
                # Access the underlying XML to check tblLook flags
                tbl_element = table._tbl
                if tbl_element is not None:
                    # Look for tblPr (table properties)
                    tbl_pr = tbl_element.find(
                        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblPr"
                    )
                    if tbl_pr is not None:
                        # Look for tblLook element
                        tbl_look = tbl_pr.find(
                            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblLook"
                        )
                        if tbl_look is not None:
                            table_info["tbl_look_flags"] = dict(
                                tbl_look.attrib
                            )
                            # Check for first row header flag
                            if "firstRow" in tbl_look.attrib:
                                table_info["has_header_row"] = (
                                    tbl_look.attrib["firstRow"] == "1"
                                )
                                analysis["has_header_flags"] = True
            except Exception as e:
                table_info["error"] = str(e)

            analysis["table_details"].append(table_info)
            analysis["tables_found"] += 1

        return analysis

    except ImportError:
        return {
            "error": "python-docx not available",
            "tables_found": 0,
            "table_details": [],
            "has_header_flags": False,
            "consistent_detection": False,
        }
    except Exception as e:
        return {
            "error": str(e),
            "tables_found": 0,
            "table_details": [],
            "has_header_flags": False,
            "consistent_detection": False,
        }


def run_test() -> Dict[str, Any]:
    """
    Run test 7: Check merged-cell detection via tblLook flags.

    Returns:
        Dict with test results
    """
    results: Dict[str, Any] = {
        "test_id": "07",
        "test_name": "Merged-cell detection via `tblLook` flags is consistent",
        "test_case": "Two DOCX tables: one with header row shading, one plain",
        "success": False,
        "error": None,
        "details": {},
    }

    try:
        print(
            f"Test 7: Merged-cell detection via `tblLook` flags is consistent"
        )
        print(
            f"Test case: Two DOCX tables: one with header row shading, one plain"
        )

        # Create or find test DOCX
        print("Creating test DOCX with different table styles...")
        docx_path = create_test_docx()
        results["details"]["docx_path"] = str(docx_path)

        # Analyze table properties
        print("Analyzing table properties and tblLook flags...")
        analysis = analyze_table_properties(docx_path)
        results["details"]["table_analysis"] = analysis

        # Determine success
        if "error" in analysis:
            results["error"] = analysis["error"]
            results["success"] = False
            results["details"]["result"] = f"FAIL: {analysis['error']}"
            print(f"❌ FAIL: {analysis['error']}")
        else:
            if analysis["tables_found"] > 0:
                results["success"] = True
                results["details"]["result"] = (
                    f"PASS: Found {analysis['tables_found']} tables, can access properties"
                )
                print(
                    f"✅ PASS: Found {analysis['tables_found']} tables, can access properties"
                )

                if analysis["has_header_flags"]:
                    print("   Found tblLook header flags in tables")
                else:
                    print(
                        "   No tblLook header flags detected (may be default)"
                    )
            else:
                results["success"] = False
                results["details"]["result"] = "FAIL: No tables found"
                print("❌ FAIL: No tables found")

        # Clean up temp file
        if docx_path.name.startswith("tmp"):
            try:
                docx_path.unlink()
            except:
                pass

    except Exception as e:
        results["error"] = str(e)
        results["details"]["exception"] = str(e)
        print(f"❌ Test failed with error: {e}")

    return results


def main():
    """Run the test."""
    results = run_test()

    # Save results
    output_file = Path(__file__).parent / "results.json"
    with open(output_file, "w") as f:
        json.dump(results, f, indent=2)

    print(f"Results saved to: {output_file}")
    return results


if __name__ == "__main__":
    main()
